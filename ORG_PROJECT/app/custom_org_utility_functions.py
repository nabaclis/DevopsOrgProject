import os, sys, json

import snapi.snapi_request
import snapi.snapi_asset
import snapi.snapi_meter
import snapi.snapi_plex
import slutils.sladmin
import snapi_base.keys as keys

from docx import Document
from docx.shared import Pt

import syslog

elastic='elastic.snaplogic.com'
clouddev='clouddev.snaplogic.com'

import app

#def get_asset_api_ptr(pod_prefix, config_file=None):
#    #---------------------------------------------------------------------------------------------------
#    #-
#    #---------------------------------------------------------------------------------------------------
#    admin_name, api_key, admin_uri = slutils.sladmin.init_sladmin(pod_prefix, config_file)
#    session = snapi.snapi_request.SnapiRequest(admin_name, api_key)
#    asset_api = snapi.snapi_asset.SnapiAsset(session, admin_uri)
#    sysadmin_context = snapi.snapi_context.SnapiContext(admin_name,
#                                                        api_key,
#                                                        admin_uri,
#                                                        sldb_uri=admin_uri)
#    sysadmin_snapi_snap_pack = sysadmin_context.schema_manager
#    return asset_api,sysadmin_snapi_snap_pack
#
#def add_subscription(org_name, sysadmin_snapi_snap_pack, snaps_dict, required_snap_list):
#    #---------------------------------------------------------------------------------------------------
#    #-
#    #---------------------------------------------------------------------------------------------------
#    subs=[]
#    print '++++++++++SNAPS-DICT:',snaps_dict
#    for key in required_snap_list:
#        snaps_dict[key]['is_sub']=True
#        subs.append(snaps_dict[key])
#    org_snid=asset_api.lookup_org(org_name)['snode_id']
#    sysadmin_snapi_snap_pack.modify_subscriptions(org_snid, {keys.SNAP_PACK_SUBSCRIPTIONS: subs})

def get_file_names():
    #---------------------------------------------------------------------------------------------------
    #-
    #---------------------------------------------------------------------------------------------------
    message_buffer=[]
    keys_file_name=os.environ['HOME']+'/snaplogic/Tectonic/etc/keys.properties'
    if not(os.path.exists(keys_file_name) and os.path.isfile(keys_file_name)):
       message_buffer.append('Error: Missing keys.properties file:'+input_file)
       syslog.syslog('Error: Missing keys.properties file:%s'%(input_file))
       return 'ERROR'
    return keys_file_name


def get_pod_prefix(pod_name, elastic,clouddev, keys_file):
    #---------------------------------------------------------------------------------------------------
    #-    Added 6/20/2017 from build_org.py
    #---------------------------------------------------------------------------------------------------
    error_messages=[]
    message_buffer=[]
    if pod_name in ['prodxl','prodxl2']:
       uri='https://'+elastic
    elif pod_name in ['canaryxl', 'canaryxl2']:
       uri='https://canary.'+elastic
    elif pod_name == 'uatxl':
       uri='https://uat.'+elastic
    elif pod_name == 'ux':
       uri='https://ux.'+clouddev
    elif pod_name == 'ux2':
       uri='https://ux2.'+clouddev
    elif pod_name == 'ux3':
       uri='https://ux3.'+clouddev
    elif pod_name == 'portal':
       uri='https://portal.'+elastic
    elif pod_name == 'qa':
       uri='https://qa.'+clouddev
    elif pod_name == 'budgy':
       uri='https://budgy.'+elastic
    elif pod_name == 'spark':
       uri='https://spark.'+elastic
    elif pod_name == 'salespod':
       uri='https://SnapLogicSales.'+elastic
    elif pod_name == 'snap':
       uri='https://snapxl.'+elastic
    elif pod_name == 'perf':
       uri='https://perfxl.'+elastic
    elif pod_name == 'prov-sldb':
       uri='https://prov-sldb.'+clouddev
    elif pod_name == 'dev-sldb':
       uri='https://dev-sldb.'+clouddev
    elif pod_name == 'stage':
       uri='https://stage.'+elastic
    else:
       uri='https://'+pod_name+'.'+elastic
    with open(keys_file) as fp:
       for i in fp:
          if i.find(uri) > 0:
             j=i.split('=')
             if len(j) > 0:
                k=j[0].split('.')
             if len(k) > 0:
                return k[0].strip(), error_messages
    message_buffer.append("""No keys exist in the keys.properties file for pod name:%s"""%(pod_name))
    syslog.syslog("""No keys exist in the keys.properties file for pod name: %s"""%(pod_name))
    error_messages.append("""No keys exist in the keys.properties file for pod name:%s"""%(pod_name))

    return None, error_messages


def get_elastic_mrc():
    #-----------------------------------------------------------------------------------------------------------------------
    #
    #-----------------------------------------------------------------------------------------------------------------------
    build_results=os.popen('curl https://elastic.snaplogic.com/status 2> /dev/null').read()
    a_json=json.loads(build_results)
    if 'build_tag' in a_json:
       return str(a_json['build_tag'])
    else:
       print 'Terminating program because the MRC could not be determined...'
       sys.exit()

def does_org_exist(pod_prefix,org_name):
    #----------------------------------------------------------------------------------------------
    #-
    #---------------------------------------------------------------------------------------------
    org_name = org_name+'/'
    try:
       admin_name, api_key, admin_uri = slutils.sladmin.init_sladmin(pod_prefix, None)
       session = snapi.snapi_request.SnapiRequest(admin_name, api_key)
       asset_api = snapi.snapi_asset.SnapiAsset(session, admin_uri)
       plex_api = snapi.snapi_plex.SnapiPlex(session, admin_uri)
       meter_api = snapi.snapi_meter.SnapiMeter(session, admin_uri)
    except:
       return False
    for i in plex_api.list_all_plexes():
        if i['runtime_path_id'].find(org_name) >= 0:
           print '>>>>>>ORG EXISTS.......'
           return True

def generate_groundplex_instructions_form(org_name,pod,keys,template_file,environment='dev'):
    #--------------------------------------------------------------------------------------
    #-
    #--------------------------------------------------------------------------------------
    document = Document(template_file)
    mrc_build=get_elastic_mrc()
    output_form_dir=app.USER_HOME+'/'+'GROUNDPLEX-FORMS-DIR'
    for i,j in enumerate(document.paragraphs):
        if j.text.find('<ORG-NAME>') >= 0:
           document.paragraphs[i].text= document.paragraphs[i].text.replace('<ORG-NAME>',org_name)
        if j.text.find('<ENV>') >= 0:
           document.paragraphs[i].text= document.paragraphs[i].text.replace('<ENV>',environment)
   
    for i,j in enumerate(document.paragraphs):
        if j.text.find('Please note: The cc.api_key') >= 0:
           style = document.styles['Normal']
           font = style.font
           font.name = 'Menlo'
           font.size = Pt(8)
           #font.bold=True
           document.paragraphs[i].insert_paragraph_before(keys)
    if not (os.path.exists(output_form_dir) and os.path.isdir(output_form_dir)):
       os.mkdir(output_form_dir)
    document.save("""%s/%s-%s.docx"""%(output_form_dir,org_name,pod))
    print """\n%s/%s-%s.docx has Groundplex keys\n"""%(output_form_dir,org_name,pod) 

def generate_property_keys(org_name,prefix):
    #--------------------------------------------------------------------------------------
    #-
    #--------------------------------------------------------------------------------------
    groundplex_template_file="""%s/app/CUSTOMER-TEMPLATE-FORM/Template-Groundplex-Instructions.docx"""%(app.APP_HOME) 
    environment='Sidekick-dev'
    if prefix.find('prod')==0:
       prefix='prod-operator'
       location='Elastic'
    elif prefix.find('canary')==0:
       prefix='canary-operator'
       location='Canary'
    elif prefix.find('uat')==0:
       prefix='uat-operator'
       location='UAT'
    elif prefix.find('salespod')==0:
       prefix='salespod-operator'
       location='Sales'
    else:
       return 'Error: Could not generate Keys because of unknown prefix'
    if does_org_exist(prefix,org_name) != True:
       print 'Error: The Org does not exist so Keys cannot be created...'
       return 'Error: The Org does not exist so Keys cannot be created...'
    else:
       syslog.syslog('Success: The Org exists so Keys can be created...')
    cmd="""%s/Tectonic/cloudops/tools/getkey.py --prefix=%s --subscriber_id=%s --target_prefix=cc"""%(app.SNAPLOGIC_HOME,prefix,org_name)
    try:
      syslog.syslog('Command:'+cmd)
      results=os.popen(cmd).read()
      generate_groundplex_instructions_form(org_name,location,results,groundplex_template_file,environment)
      return results
    except:
      syslog.syslog('Failed command:'+cmd)
      return None

